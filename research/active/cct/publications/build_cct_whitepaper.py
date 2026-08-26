from __future__ import annotations

import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
HERE = Path(__file__).resolve().parent
OUT = ROOT / "output" / "docx" / "CCT-livre-blanc.docx"
FIGURE = HERE / "cct-architecture.png"

NAVY = "17324D"
TEAL = "177E89"
GOLD = "B58B2A"
INK = "24323D"
MUTED = "66737D"
PALE = "EEF5F4"
LIGHT = "F3F5F6"
WHITE = "FFFFFF"
TABLE_WIDTH = 9360
BREAK_CHAPTERS: set[int] = set()

TOC_GROUPS = (
    (
        "Fondements et décision",
        (
            "Résumé et méthode de lecture",
            "1. Une politique des interdépendances",
            "2. Principes constitutionnels non négociables",
            "3. Les unités politiques",
            "4. Répartition des compétences",
            "5. Institutions démocratiques",
            "6. Comment une décision est prise",
        ),
    ),
    (
        "Économie, écologie et garanties",
        (
            "7. Constitution économique",
            "8. Métabolisme écologique",
            "9. Justice globale et réparation",
            "10. Information, langues, connaissance et numérique",
            "11. Justice, sûreté et force",
            "12. Urgence sans dictature provisoire permanente",
            "13. Où se cachent les centres de pouvoir",
            "14. Prévenir la reconstitution des classes dominantes",
            "15. Métriques et adaptation stratégique",
            "16. Cycle de maintenance constitutionnelle",
        ),
    ),
    (
        "Transition, épreuves et reconstruction",
        (
            "17. Stratégie de transition",
            "18. Tests de résistance",
            "19. Tableau de bord minimal",
            "20. Décisions politiques encore ouvertes",
            "21. Conditions d’échec du modèle",
            "22. Formule constitutionnelle courte",
            "23. Conclusion et conditions de révision",
            "Annexes — résultats et carte de validation",
        ),
    ),
)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run(run, *, size=None, color=INK, bold=None, italic=None, font="Aptos") -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_inline(paragraph, text: str) -> None:
    pattern = re.compile(r"(\*\*.*?\*\*|\*.*?\*)")
    for part in pattern.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2]); set_run(run, bold=True)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1]); set_run(run, italic=True)
        else:
            run = paragraph.add_run(part); set_run(run)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page "); set_run(run, size=8.5, color=MUTED)
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def new_numbering_id(doc: Document) -> int:
    """Create a fresh level-0 decimal sequence so each Markdown list restarts at 1."""
    numbering = doc.part.numbering_part.element
    abstract_id = "0"
    for abstract in numbering.findall(qn("w:abstractNum")):
        styles = abstract.findall(f".//{qn('w:pStyle')}")
        if any(style.get(qn("w:val")) == "ListNumber" for style in styles):
            abstract_id = abstract.get(qn("w:abstractNumId"))
            break
    existing = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(existing, default=0) + 1
    num = OxmlElement("w:num"); num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId"); abstract.set(qn("w:val"), abstract_id); num.append(abstract)
    override = OxmlElement("w:lvlOverride"); override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride"); start.set(qn("w:val"), "1"); override.append(start)
    num.append(override); numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr"); p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId"); num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl); num_pr.append(num)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5); section.page_height = Inches(11)
    section.top_margin = Inches(0.82); section.bottom_margin = Inches(1.03)
    section.left_margin = Inches(0.9); section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35); section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"; normal.font.size = Pt(10.7); normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(7); normal.paragraph_format.line_spacing = 1.22
    for name, size, color, before, after in (
        ("Title", 31, NAVY, 0, 10), ("Subtitle", 14, TEAL, 0, 10),
        ("Heading 1", 19, NAVY, 18, 9), ("Heading 2", 14, TEAL, 14, 7),
        ("Heading 3", 11.5, NAVY, 10, 4),
    ):
        style = styles[name]
        style.font.name = "Aptos Display" if name != "Heading 3" else "Aptos"
        style.font.size = Pt(size); style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Aptos"; style.font.size = Pt(10.7)
        style.paragraph_format.left_indent = Inches(0.42)
        style.paragraph_format.first_line_indent = Inches(-0.20)
        style.paragraph_format.space_after = Pt(4); style.paragraph_format.line_spacing = 1.18
    quote = styles["Quote"]
    quote.font.name = "Aptos"; quote.font.size = Pt(11.2); quote.font.italic = False
    quote.font.color.rgb = RGBColor.from_string(NAVY)
    quote.paragraph_format.left_indent = Inches(0.28); quote.paragraph_format.right_indent = Inches(0.25)
    quote.paragraph_format.space_before = Pt(8); quote.paragraph_format.space_after = Pt(10)

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "CONFÉDÉRATION DES COMMUNS TERRESTRES   /   LIVRE BLANC"
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(hp.runs[0], size=8.1, color=MUTED, bold=True)
    add_page_field(section.footer.paragraphs[0])


def make_figure() -> None:
    image = Image.new("RGB", (1692, 828), "white")
    draw = ImageDraw.Draw(image)
    regular_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
    bold_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"
    regular = ImageFont.truetype(regular_path, 31)
    small = ImageFont.truetype(regular_path, 25)
    bold = ImageFont.truetype(bold_path, 30)
    title_font = ImageFont.truetype(bold_path, 40)
    footer_font = ImageFont.truetype(bold_path, 25)
    title = "Une architecture polycentrique, trois pouvoirs distincts"
    box = draw.textbbox((0, 0), title, font=title_font)
    draw.text(((1692 - (box[2] - box[0])) / 2, 70), title, font=title_font, fill="#17324D")
    boxes = [
        (70, 280, 490, 560, "Territoires démocratiques", "Citoyenneté, droits,\nservices et proximité"),
        (635, 280, 1055, 560, "Communs fonctionnels", "Ressources et réseaux\nselon leurs bassins réels"),
        (1200, 280, 1620, 560, "Confédération planétaire", "Droits, paix, communs\nmondiaux et redistribution"),
    ]
    for x1, y1, x2, y2, heading, body in boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=24, fill="#EEF5F4", outline="#177E89", width=4)
        hbox = draw.textbbox((0, 0), heading, font=bold)
        draw.text(((x1+x2-(hbox[2]-hbox[0]))/2, y1+62), heading, font=bold, fill="#17324D")
        lines = body.split("\n")
        for idx, line in enumerate(lines):
            bbox = draw.textbbox((0, 0), line, font=small)
            draw.text(((x1+x2-(bbox[2]-bbox[0]))/2, y1+145+idx*38), line, font=small, fill="#44525C")
    for start, end in (((500, 420), (625, 420)), ((1065, 420), (1190, 420))):
        draw.line((start, end), fill="#B58B2A", width=5)
        draw.polygon([(start[0], start[1]), (start[0]+18, start[1]-10), (start[0]+18, start[1]+10)], fill="#B58B2A")
        draw.polygon([(end[0], end[1]), (end[0]-18, end[1]-10), (end[0]-18, end[1]+10)], fill="#B58B2A")
    footer = "Garantie commune  •  Exécution distribuée  •  Contrôle croisé  •  Révocation possible"
    bbox = draw.textbbox((0, 0), footer, font=footer_font)
    draw.text(((1692-(bbox[2]-bbox[0]))/2, 690), footer, font=footer_font, fill="#177E89")
    image.save(FIGURE)


def add_cover(doc: Document) -> None:
    for _ in range(5):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(14)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("LIVRE BLANC CONSTITUTIONNEL ET EXPÉRIMENTAL"); set_run(r, size=9.5, color=GOLD, bold=True)
    p.paragraph_format.space_after = Pt(24)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Confédération des\ncommuns terrestres"); set_run(r, size=32, color=NAVY, bold=True, font="Aptos Display")
    p.paragraph_format.space_after = Pt(15)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Une architecture écosocialiste libertaire pour coordonner le monde sans souverain mondial illimité"); set_run(r, size=14, color=TEAL, italic=True)
    p.paragraph_format.space_after = Pt(42)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ÉTAT DE RECHERCHE  •  AOÛT 2026"); set_run(r, size=11, color=MUTED, bold=True)
    p.paragraph_format.space_after = Pt(60)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PROPOSITION CONSOLIDÉE — NON VALIDÉE EMPIRIQUEMENT"); set_run(r, size=9, color=WHITE, bold=True)
    p_pr = p._p.get_or_add_pPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), NAVY); p_pr.append(shd)
    p.paragraph_format.left_indent = Inches(1.15); p.paragraph_format.right_indent = Inches(1.15)
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(8)
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph("Sommaire", style="Heading 1")
    p.paragraph_format.page_break_before = False
    for group, entries in TOC_GROUPS:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(7); p.paragraph_format.space_after = Pt(2)
        run = p.add_run(group.upper()); set_run(run, size=8.5, color=TEAL, bold=True)
        for entry in entries:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.space_after = Pt(1.5); p.paragraph_format.line_spacing = 1.0
            run = p.add_run(entry); set_run(run, size=8.8, color=INK)
    note = doc.add_paragraph("Ce document distingue les garanties constitutionnelles, les mécanismes proposés et les résultats expérimentaux. Les annexes conservent les échecs qui ont modifié le modèle.")
    note.style = doc.styles["Quote"]
    note.paragraph_format.line_spacing = 1.0
    note.paragraph_format.space_before = Pt(5); note.paragraph_format.space_after = Pt(4)
    for run in note.runs:
        set_run(run, size=8.8, color=NAVY)
    doc.add_page_break()


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = len(rows[0])
    widths = [int(TABLE_WIDTH / cols)] * cols
    widths[-1] += TABLE_WIDTH - sum(widths)
    if cols == 2:
        widths = [2700, 6660]
    elif cols == 3:
        widths = [2200, 3000, 4160]
    elif cols == 5:
        widths = [3000, 1270, 1270, 1270, 2550]
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for r_index, row in enumerate(rows):
        for c_index, text in enumerate(row):
            cell = table.cell(r_index, c_index)
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline(p, text)
            p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.08
            for run in p.runs:
                set_run(run, size=8.8, color=WHITE if r_index == 0 else INK, bold=(r_index == 0))
            if r_index == 0:
                set_cell_shading(cell, NAVY)
            elif r_index % 2 == 0:
                set_cell_shading(cell, LIGHT)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, widths)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)


def add_markdown(doc: Document, text: str, *, strip_model_front=False) -> None:
    lines = text.splitlines()
    if strip_model_front:
        start = next(i for i, line in enumerate(lines) if line.startswith("## 1. Conclusion architecturale"))
        lines = lines[start:]
    paragraph_buffer: list[str] = []
    table_rows: list[list[str]] = []
    in_code = False
    diagram_inserted = False
    current_num_id: int | None = None

    def end_numbered_list() -> None:
        nonlocal current_num_id
        current_num_id = None

    def flush_para() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            p = doc.add_paragraph()
            add_inline(p, " ".join(paragraph_buffer))
            paragraph_buffer = []

    def flush_table() -> None:
        nonlocal table_rows
        if table_rows:
            rows = [row for row in table_rows if not all(re.fullmatch(r"[-: ]+", cell) for cell in row)]
            add_table(doc, rows); table_rows = []

    for raw in lines:
        line = raw.strip()
        if line.startswith("```"):
            flush_para(); flush_table(); end_numbered_list()
            in_code = not in_code
            if in_code and "mermaid" in line and not diagram_inserted:
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run()
                figure = r.add_picture(str(FIGURE), width=Inches(5.55))
                figure._inline.docPr.set(
                    "descr",
                    "Architecture polycentrique de la Confédération des communs terrestres : "
                    "articulation des territoires, communs fonctionnels et institutions mondiales bornées.",
                )
                figure._inline.docPr.set("title", "Architecture institutionnelle de la CCT")
                caption = doc.add_paragraph("Figure 1 — Architecture institutionnelle de la CCT")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in caption.runs: set_run(run, size=8.5, color=MUTED, italic=True)
                diagram_inserted = True
            continue
        if in_code:
            continue
        if line.startswith("|") and line.endswith("|"):
            flush_para(); end_numbered_list()
            table_rows.append([cell.strip() for cell in line.strip("|").split("|")])
            continue
        flush_table()
        if not line:
            flush_para(); end_numbered_list(); continue
        heading = re.match(r"^(#{1,4})\s+(.*)$", line)
        if heading:
            flush_para(); end_numbered_list()
            level = len(heading.group(1)); title = heading.group(2)
            style_level = 1 if level <= 2 else min(3, level - 1)
            p = doc.add_paragraph(title, style=f"Heading {style_level}")
            chapter = re.match(r"(\d+)\.", title)
            if style_level == 1 and chapter and int(chapter.group(1)) in BREAK_CHAPTERS:
                p.paragraph_format.page_break_before = True
            if style_level == 1 and title.startswith("Annexe A"):
                p.paragraph_format.page_break_before = True
            continue
        if line.startswith(">"):
            flush_para(); end_numbered_list()
            p = doc.add_paragraph(style="Quote"); add_inline(p, line.lstrip("> "))
            p_pr = p._p.get_or_add_pPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), PALE); p_pr.append(shd)
            continue
        if re.match(r"^-\s+", line):
            flush_para(); end_numbered_list(); p = doc.add_paragraph(style="List Bullet"); add_inline(p, re.sub(r"^-\s+", "", line)); continue
        if re.match(r"^\d+\.\s+", line):
            flush_para()
            if current_num_id is None:
                current_num_id = new_numbering_id(doc)
            p = doc.add_paragraph(style="List Number"); apply_numbering(p, current_num_id)
            add_inline(p, re.sub(r"^\d+\.\s+", "", line)); continue
        end_numbered_list()
        paragraph_buffer.append(line)
    flush_para(); flush_table()


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    make_figure()
    doc = Document(); configure_document(doc); add_cover(doc); add_toc(doc)
    short_text = (HERE / "livre-blanc-cct.md").read_text(encoding="utf-8")
    add_markdown(doc, short_text.split("## 1. Une politique des interdépendances", 1)[0])
    model_text = (ROOT / "modele-gouvernance-ecosocialiste-libertaire.md").read_text(encoding="utf-8")
    model_text = model_text[model_text.index("## 1. Conclusion architecturale"):model_text.index("## 24. Statut expérimental")]
    model_text = model_text.replace(
        "## 18. Tests de résistance",
        """### Ce que les essais ont déjà changé

Les essais synthétiques ont révélé que des protections solides isolément peuvent échouer ensemble lorsqu’elles partagent les mêmes réseaux, experts, clés, fournisseurs ou réserves. Ils ont aussi conservé trois erreurs de méthode : un avantage initial indu, des tirages non appariés et une mesure d’empreinte erronée. Les mécanismes retenus en portent la trace : voies hors ligne pour les droits vitaux, règle publique de pénurie, registre des dépendances, budget de charge, extinction séparée des pouvoirs temporaires et test de polycrise.

Les prototypes exécutables ont confirmé une limite plus étroite : une règle ne rend pas une action possible, un gain matériel ne devient pas une capacité vérifiée sans attestation indépendante, et une réparation peut exiger plusieurs étapes. Une restauration séquencée concorde avec un oracle indépendant dans une abstraction finie ; elle ne crée ni les ressources, ni les témoins, ni les institutions réelles dont dépendrait son usage. Son statut reste local, non promouvable et sans transport externe établi.

## 18. Tests de résistance""",
    )
    model_text = model_text.replace(
        "## 16. Cycle de maintenance constitutionnelle",
        """### Charge constitutionnelle et polycrise

La CCT doit être éprouvée lorsque besoins vitaux, plafond écologique, droits, trace et restitution sollicitent les mêmes ressources dégradées. Les rapports secondaires et les formalités réversibles peuvent alors être délestés ; accès vital, traçabilité minimale, recours contre coercition, plafond critique et preuve de restitution ne le peuvent pas. Une procédure qui n’ajoute aucune protection observable face à une solution plus simple doit être fusionnée ou retirée.

## 16. Cycle de maintenance constitutionnelle""",
    )
    model_text = model_text.replace(
        "## 23. Conclusion",
        """### Conditions de révision

La proposition doit être reconstruite ou retirée si des observations indépendantes montrent durablement que ses collèges paralysent les crises, que ses communs deviennent des corporations, que les droits portables échouent à protéger la sortie, que les urgences conservent leurs pouvoirs ou que la complexité impose davantage de dépendance qu’elle ne produit de contrôle. Aucun test local ni aucun texte ne tranche ces questions : ils fixent seulement les conditions dans lesquelles une réponse future devrait compter.

## 23. Conclusion""",
    )
    add_markdown(doc, model_text)
    add_markdown(doc, """# Annexe — Carte de validation\n\n| Élément | Écrit | Logiciel | Synthétique | Structurel borné | Terrain / réobservation |\n|---|---:|---:|---:|---:|---:|\n| Architecture politique complète | Oui | Partiel | Non | Non | Non |\n| Continuité et planification | Oui | Oui | Oui | Non | Non |\n| Droits portables et recours | Oui | Partiel | Non | Non | Non |\n| Pouvoirs temporaires et restitution | Oui | Oui | Partiel | Non | Non |\n| Restauration séquencée | Oui | Oui | Non | Oui | Non |\n\nCes colonnes ne se compensent pas : un test logiciel établit une exécution définie, une simulation établit un résultat dans son monde, et une vérification structurelle établit une propriété de son abstraction. Aucune ne valide à elle seule un effet territorial ni une capacité institutionnelle générale.""")
    props = doc.core_properties
    props.title = "Confédération des communs terrestres — Livre blanc"
    props.subject = "Architecture écosocialiste libertaire, modèle institutionnel et programme expérimental"
    props.author = "Confédération des communs terrestres — document de travail"
    props.keywords = "CCT, écosocialisme, libertaire, gouvernance mondiale, communs, constitution"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
