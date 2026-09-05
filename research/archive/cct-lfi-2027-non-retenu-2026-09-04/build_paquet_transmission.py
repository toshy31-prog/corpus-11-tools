from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "output" / "CCT-France-paquet-transmission-LFI-2027.docx"
TABLE_W = 9360
NAVY, TEAL, INK, MUTED, PALE, LIGHT = "17324D", "177E89", "24323D", "66737D", "EAF1F4", "F3F5F6"


def run_style(run, size=10.7, color=INK, bold=False, italic=False, font="Aptos"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold, run.italic = bold, italic


def shade(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    node = tcpr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd"); tcpr.append(node)
    node.set(qn("w:fill"), fill)


def margins(cell):
    tcpr = cell._tc.get_or_add_tcPr()
    mar = tcpr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar"); tcpr.append(mar)
    for name, value in (("top", 90), ("bottom", 90), ("start", 120), ("end", 120)):
        n = mar.find(qn(f"w:{name}"))
        if n is None:
            n = OxmlElement(f"w:{name}"); mar.append(n)
        n.set(qn("w:w"), str(value)); n.set(qn("w:type"), "dxa")


def table_geometry(table, widths):
    table.autofit = False; table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblpr = table._tbl.tblPr
    for tag, value in (("tblW", sum(widths)), ("tblInd", 120)):
        node = tblpr.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}"); tblpr.append(node)
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for node in list(grid): grid.remove(node)
    for width in widths:
        node = OxmlElement("w:gridCol"); node.set(qn("w:w"), str(width)); grid.append(node)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths[i] / 1440)
            tcw = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tcw.set(qn("w:w"), str(widths[i])); tcw.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER; margins(cell)


def set_repeat_header(row):
    p = row._tr.get_or_add_trPr(); tag = OxmlElement("w:tblHeader"); tag.set(qn("w:val"), "true"); p.append(tag)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = paragraph.add_run("CCT France  |  "); run_style(r, 8.2, MUTED, bold=True)
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE"); paragraph._p.append(fld)


def setup(doc):
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.top_margin, sec.bottom_margin = Inches(.82), Inches(.85)
    sec.left_margin, sec.right_margin = Inches(.9), Inches(.9)
    sec.header_distance, sec.footer_distance = Inches(.32), Inches(.32)
    styles = doc.styles
    normal = styles["Normal"]; normal.font.name = "Aptos"; normal.font.size = Pt(10.7); normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.15
    for name, size, color, before, after in (("Heading 1", 18, NAVY, 17, 8), ("Heading 2", 13.5, TEAL, 12, 6), ("Heading 3", 11.5, NAVY, 8, 4)):
        style = styles[name]; style.font.name = "Aptos Display"; style.font.size = Pt(size); style.font.bold = True; style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after); style.paragraph_format.keep_with_next = True
    header = sec.header.paragraphs[0]; header.text = "CCT FRANCE  /  PAQUET DE TRANSMISSION"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT; run_style(header.runs[0], 8.2, MUTED, bold=True)
    add_page_number(sec.footer.paragraphs[0])


def para(doc, text, *, style=None, size=10.7, color=INK, bold=False, italic=False, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None: p.alignment = align
    run = p.add_run(text); run_style(run, size, color, bold, italic)
    return p


def labelled(doc, label, value):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(label + " "); run_style(r, 10.3, NAVY, bold=True)
    r = p.add_run(value); run_style(r, 10.3)
    return p


def source_sections():
    text = (ROOT / "03-mesures" / "catalogue.md").read_text(encoding="utf-8")
    out = []
    for n in range(1, 16):
        token = f"## M{n:02d}"
        part = text.split(token, 1)[1].split("\n## ", 1)[0].strip()
        title, body = part.split("\n\n", 1)
        fields = {}
        for label in ("Effet", "Porteur", "Voie", "Risque", "Garde-fou", "Indicateurs", "Retrait"):
            match = re.search(rf"\*\*{re.escape(label)}\.\*\* (.*?)(?= \*\*[A-ZÉ][^*]*\.\*\*|$)", body, flags=re.S)
            if match: fields[label] = " ".join(match.group(1).split())
        out.append((f"M{n:02d}", title.strip(" —"), fields))
    return out


def main():
    OUT.parent.mkdir(exist_ok=True)
    doc = Document(); setup(doc)

    para(doc, "CONTRIBUTION PROGRAMMATIQUE · LFI 2027", size=10, color=TEAL, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Planifier sans confisquer le pouvoir", size=29, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Six garanties distinctives pour L'Avenir en commun", size=14, color=TEAL, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "France · septembre 2026 · paquet de transmission", size=9.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    t = doc.add_table(rows=1, cols=1); table_geometry(t, [TABLE_W]); cell = t.cell(0, 0); shade(cell, PALE)
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Ce qui est proposé"); run_style(r, 11, NAVY, bold=True)
    p = cell.add_paragraph(); p.paragraph_format.space_after = Pt(1)
    r = p.add_run("Six garanties transversales à intégrer, modifier ou écarter séparément ; neuf fiches de mise en oeuvre déjà ancrées dans le socle LFI."); run_style(r, 10.5)
    para(doc, "Demande", style="Heading 1")
    para(doc, "Nous vous transmettons, en 2026, six garanties proposées à L'Avenir en commun. Elles complètent des engagements déjà présents dans l'édition 2025 actuellement mobilisée pour le programme 2027 : lois-cadres de planification écologique et démocratique, Conseil à la planification écologique, moyens des opérateurs publics, protection des biens communs, intervention citoyenne et droits nouveaux des salariés.")
    para(doc, "Le dossier ne demande pas l'adoption d'une architecture CCT complète, ni ne présente le socle LFI comme une invention CCT. Ses six apports distinctifs portent sur les dépendances vitales, la capacité de continuité, l'accès hors numérique, l'extinction des exceptions, la charge démocratique et l'évaluation contradictoire.")
    para(doc, "Statut de travail", style="Heading 1")
    para(doc, "Le paquet a été constitué sans intervenant extérieur avant envoi. Il ne prétend donc ni avoir été reçu, ni validé juridiquement, ni chiffré définitivement. Les points ouverts sont explicitement signalés dans chaque fiche ; ils ne sont pas masqués par une promesse.")

    para(doc, "Les six garanties à examiner en priorité", style="Heading 1")
    para(doc, "L'audit de non-duplication sépare les apports réellement distinctifs (ci-dessous) des neuf clauses qui ne font que préciser des engagements déjà portés par LFI.", color=MUTED)
    doc.add_page_break()
    table = doc.add_table(rows=1, cols=3); table.style = "Table Grid"; table_geometry(table, [850, 3500, 5010]); set_repeat_header(table.rows[0])
    for cell, text in zip(table.rows[0].cells, ("ID", "Amendement", "Effet recherché")):
        shade(cell, LIGHT); p = cell.paragraphs[0]; r = p.add_run(text); run_style(r, 9.2, NAVY, bold=True)
    measures = source_sections()
    priority_ids = {"M04", "M05", "M07", "M09", "M10", "M12"}
    priority_measures = [m for m in measures if m[0] in priority_ids]
    for mid, title, fields in priority_measures:
        row = table.add_row().cells
        for cell in row: cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        r = row[0].paragraphs[0].add_run(mid); run_style(r, 9.2, NAVY, bold=True)
        r = row[1].paragraphs[0].add_run(title); run_style(r, 9.2, INK, bold=True)
        r = row[2].paragraphs[0].add_run(fields.get("Effet", "")); run_style(r, 9.1)

    doc.add_page_break()
    para(doc, "Fiches des six garanties distinctives", style="Heading 1")
    para(doc, "Chaque fiche indique l'effet, la voie, le risque et le garde-fou. Les conditions de retrait constituent une discipline : elles empêchent de défendre un dispositif indépendamment de ses effets.", color=MUTED)
    for i, (mid, title, fields) in enumerate(priority_measures):
        para(doc, f"{mid} — {title}", style="Heading 2")
        labelled(doc, "Effet.", fields.get("Effet", "À préciser."))
        labelled(doc, "Voie.", fields.get("Voie", "À qualifier."))
        labelled(doc, "Risque.", fields.get("Risque", "À qualifier."))
        labelled(doc, "Garde-fou.", fields.get("Garde-fou", "À préciser."))
        labelled(doc, "Indicateurs.", fields.get("Indicateurs", "À préciser."))
        if "Retrait" in fields: labelled(doc, "Retrait.", fields["Retrait"])
        if i == 2: doc.add_page_break()

    doc.add_page_break()
    para(doc, "Clauses de mise en oeuvre déjà ancrées dans le socle LFI", style="Heading 1")
    para(doc, "Les neuf fiches suivantes ne sont pas présentées comme des nouveautés. Elles précisent les garanties de pluralisme, continuité, recours et contrôle attachées aux engagements LFI existants.", color=MUTED)
    for mid, title, fields in measures:
        if mid in priority_ids:
            continue
        para(doc, f"{mid} — {title}", style="Heading 2")
        labelled(doc, "Effet.", fields.get("Effet", "À préciser."))
        labelled(doc, "Garde-fou.", fields.get("Garde-fou", "À préciser."))

    doc.add_page_break()
    para(doc, "Mode d'intégration dans L'Avenir en commun", style="Heading 1")
    para(doc, "La contribution est datée de 2026 et s'appuie sur les points d'ancrage de l'édition 2025 du programme, alors mobilisée pour 2027. Les formulations ci-dessous sont des ajouts programmatiques, non des articles de loi et non des citations du programme.")
    anchors = [
        ("M01–M02, M04–M05, M10–M13", "Chapitre 12 : planification écologique", "Rendre les lois-cadres, le Conseil, les données, les investissements et les évaluations contestables et continus."),
        ("M03, M06, M08", "Chapitres 2 et 14 : biens communs", "Donner aux communs de l'eau et aux pôles publics des règles de gouvernance, de recours et de continuité."),
        ("M07, M09, M15", "Chapitre 1 : intervention populaire", "Garantir l'accès effectif aux droits, l'arrêt des exceptions et le suivi citoyen des décisions."),
        ("M14", "Chapitre 3 : citoyenneté dans l'entreprise", "Prolonger les droits de préemption et de contrôle des salariés vers les infrastructures vitales abandonnées."),
    ]
    table = doc.add_table(rows=1, cols=3); table.style = "Table Grid"; table_geometry(table, [1850, 3000, 4510]); set_repeat_header(table.rows[0])
    for cell, text in zip(table.rows[0].cells, ("Mesures", "Ancrage LFI 2027", "Ajout proposé")):
        shade(cell, LIGHT); r = cell.paragraphs[0].add_run(text); run_style(r, 9.2, NAVY, bold=True)
    for values in anchors:
        row = table.add_row().cells
        for cell, value in zip(row, values): r = cell.paragraphs[0].add_run(value); run_style(r, 9.2)

    para(doc, "Limites et sources", style="Heading 1")
    para(doc, "Les ancrages de la contribution LFI 2027 reposent sur l'édition 2025 de L'Avenir en commun et sur les pistes d'actualisation publiées en 2026 : « Une République permettant l'intervention populaire », « Collectiviser les biens communs fondamentaux », « Reconnaître la citoyenneté dans l'entreprise et des droits nouveaux aux salariés », « La bifurcation écologique pour une société de l'harmonie », la démocratie écologique communale et les éco-régions. Les liens, le droit applicable, les données budgétaires et les limites sont conservés dans le dossier source.")
    para(doc, "Le document de référence est le dossier local CCT France, notamment 02b-matrice-amendements.md, 04-droit-et-competences.md, 05-budget-et-capacites.md, 08-contradiction.md et 10-auto-contradiction.md.", size=9.3, color=MUTED, italic=True)

    doc.add_page_break()
    para(doc, "Instruction juridique : ce qui doit être qualifié", style="Heading 1")
    para(doc, "Les propositions ci-dessous s'appuient sur du droit existant. Elles ne sont pas des avis juridiques et ne prétendent pas que le texte d'application est déjà écrit.", color=MUTED)
    legal = [
        ("M01, M02, M10, M12, M15", "Loi, programmation et règlements des assemblées", "Constitution art. 34 ; LOLF ; règles de participation. Définir précisément compétence, indicateurs et portée de la réponse publique."),
        ("M03", "Code de l'environnement et gouvernance des bassins", "L211-1 fournit un socle de gestion durable ; articulation avec les instances existantes à instruire."),
        ("M04, M05, M06", "Lois sectorielles et statuts des opérateurs", "Définir service vital, niveau de publicité, accès aux données, responsabilités et continuité."),
        ("M07, M08", "Droits des usagers et recours", "Le CRPA et le Défenseur des droits donnent des points d'appui ; voie hors ligne et éventuelle extension organique à qualifier."),
        ("M09", "Régimes d'urgence", "Loi pour les régimes sectoriels ; articles 16 et 36 de la Constitution à traiter séparément."),
        ("M11, M14", "Commande publique, droit du travail et procédures collectives", "Critères liés à l'objet et proportionnés ; recherche de repreneur existante mais pas de reprise automatique."),
        ("M13", "Loi de finances et contrats territoriaux", "Crédits, recettes, règles de répartition et contrôle parlementaire doivent être votés."),
    ]
    table = doc.add_table(rows=1, cols=3); table.style = "Table Grid"; table_geometry(table, [1900, 2700, 4760]); set_repeat_header(table.rows[0])
    for cell, text in zip(table.rows[0].cells, ("Mesures", "Voie", "Question non résolue à l'envoi")):
        shade(cell, LIGHT); r = cell.paragraphs[0].add_run(text); run_style(r, 9.1, NAVY, bold=True)
    for values in legal:
        row = table.add_row().cells
        for cell, value in zip(row, values): r = cell.paragraphs[0].add_run(value); run_style(r, 8.9)

    para(doc, "Budget et capacités : ordre de réalité", style="Heading 1")
    para(doc, "Aucun total n'est affiché : additionner ces mesures créerait un chiffre faux. Le coût doit être instruit par périmètre, effectif, système, investissement, maintenance, financement et hypothèses. Les lois de finances imposent une présentation sincère des charges prévisibles.")
    budget = [
        ("Architecture démocratique", "M01, M02, M10, M12, M15", "secrétariats, données, indemnisation, contre-expertise", "participants, fréquence, méthodes et systèmes définis"),
        ("Eau et continuité", "M03, M04, M05, M13", "cartographies, exercices, stocks, maintenance, réparation", "service, territoire, responsable et solution de secours identifiés"),
        ("Droits effectifs", "M07, M08, M09", "guichets, médiation, recours, formation", "droit vital, voie humaine et délai de correction fixés"),
        ("Production et achat", "M06, M11, M14", "formation, audits, ingénierie de reprise, maintenance", "diagnostic de capacité et compatibilité européenne documentés"),
    ]
    table = doc.add_table(rows=1, cols=4); table.style = "Table Grid"; table_geometry(table, [1750, 1650, 3000, 2960]); set_repeat_header(table.rows[0])
    for cell, text in zip(table.rows[0].cells, ("Bloc", "Mesures", "Dépense à isoler", "Passage à l'étape suivante")):
        shade(cell, LIGHT); r = cell.paragraphs[0].add_run(text); run_style(r, 8.8, NAVY, bold=True)
    for values in budget:
        row = table.add_row().cells
        for cell, value in zip(row, values): r = cell.paragraphs[0].add_run(value); run_style(r, 8.6)
    para(doc, "Contrôle de sincérité : aucune aide sélective ou reprise ne doit être présentée comme automatiquement compatible avec le droit de l'Union ; aucune voie hors ligne comme financée sans effectifs et délais ; aucun fonds comme réel sans ressource et règles de décaissement.", size=9.4, color=MUTED, italic=True)
    para(doc, "Objections et réponse courte", style="Heading 1")
    objections = [
        ("Une bureaucratie de plus", "Chaque organe doit avoir droit, budget, trace et clause de retrait ; les formalités inutiles doivent disparaître."),
        ("Un frein à l'urgence", "Les procédures conservatoires et les délais courts préservent l'action ; l'exception sans fin devient elle-même un risque."),
        ("Une capture des communs", "Représentation pluraliste, données publiques, recours et moyens de participation limitent la capture sans prétendre l'abolir."),
        ("Un coût non financé", "Aucun montant total n'est inventé : chaque crédit doit indiquer sa recette, ses métiers et ses coûts de maintenance."),
        ("Une incompatibilité européenne", "La commande publique doit rester liée à l'objet, transparente et proportionnée ; les aides sélectives sont qualifiées avant engagement."),
    ]
    for title, answer in objections:
        labelled(doc, title + ".", answer)
    para(doc, "Ce que le paquet demande", style="Heading 1")
    para(doc, "Pas l'adoption d'un bloc fermé. Il propose six garanties prioritaires, chacune révisable ou écartable, et neuf clauses de mise en oeuvre. Il demande que les engagements déjà présents soient dotés de règles de contrôle, de recours, de continuité et de retrait à la hauteur de leurs ambitions.")
    para(doc, "État au jour de l'envoi : propositions rédigées, sourcées et auto-contradictoires ; ni réception, ni validation juridique, ni chiffrage définitif ne sont présumés.", size=9.4, color=MUTED, italic=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
