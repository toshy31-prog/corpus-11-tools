from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).parent
OUT = ROOT / "output" / "20-note-politique-envoi.docx"
NAVY = "173B5C"; TEAL = "168C9B"; INK = "29333D"; MUTED = "687787"; PALE = "EAF2F5"

def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), color); tc_pr.append(shd)

def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc; tc_pr = tc.get_or_add_tcPr(); margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None: margins = OxmlElement("w:tcMar"); tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None: node = OxmlElement(f"w:{side}"); margins.append(node)
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")

def style(run, size=11, color=INK, bold=False, italic=False):
    run.font.name = "Aptos"; run._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos"); run.font.size = Pt(size); run.font.color.rgb = RGBColor.from_string(color); run.bold = bold; run.italic = italic

def para(doc, text="", size=11, color=INK, bold=False, italic=False, style_name=None, after=7, before=0, align=None):
    p = doc.add_paragraph(style=style_name); p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after); p.paragraph_format.line_spacing = 1.12
    if align is not None: p.alignment = align
    r = p.add_run(text); style(r, size, color, bold, italic)
    return p

def main():
    doc = Document(); sec = doc.sections[0]
    sec.top_margin = Inches(.75); sec.bottom_margin = Inches(.72); sec.left_margin = Inches(.85); sec.right_margin = Inches(.85)
    normal = doc.styles["Normal"]; normal.font.name = "Aptos"; normal.font.size = Pt(10.8); normal.font.color.rgb = RGBColor.from_string(INK)
    for name, size, color in (("Heading 1", 16, NAVY), ("Heading 2", 12, TEAL)):
        st = doc.styles[name]; st.font.name = "Aptos Display"; st.font.size = Pt(size); st.font.color.rgb = RGBColor.from_string(color); st.font.bold = True
        st.paragraph_format.space_before = Pt(12); st.paragraph_format.space_after = Pt(5)

    para(doc, "CONTRIBUTION PROGRAMMATIQUE · LFI 2027", 9.5, TEAL, True, after=4)
    para(doc, "Trois engagements que la bifurcation doit prendre", 21.5, NAVY, True, after=3)
    para(doc, "Des ajouts distincts pour le programme LFI 2027", 12, MUTED, after=14)

    box = doc.add_table(rows=1, cols=1); cell = box.cell(0, 0); shade(cell, PALE); set_cell_margins(cell)
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2); r = p.add_run("Le point de départ"); style(r, 10.5, NAVY, True)
    p = cell.add_paragraph(); p.paragraph_format.space_after = Pt(1); r = p.add_run("La règle verte, les éco-régions et les services publics sont déjà portés. Ces propositions ne les reformulent pas : elles répondent à trois situations où la bifurcation doit protéger concrètement la population."); style(r, 10.5)

    para(doc, "Trois ajouts à instruire", style_name="Heading 1")
    items = [
        ("Les besoins vitaux avant les privilèges.", "En cas de contrainte écologique, une règle publique de priorité protège l'eau, l'alimentation, la santé, le logement, l'énergie et la mobilité indispensables. Ni le prix seul, ni la file d'attente, ni des dérogations privées ne décident de l'accès aux besoins fondamentaux."),
        ("Pas de transfert d'un service essentiel à l'aveugle.", "Avant toute nationalisation, délégation, reprise publique ou changement d'opérateur : preuve de continuité sur les personnels, contrats, stocks, données, systèmes, accès et solutions de secours. Sans elle, le transfert est suspendu, hors mesure temporaire contrôlée."),
        ("Aucun pouvoir d'urgence ne se prolonge lui-même.", "Durée courte, autorité distincte pour toute prolongation, contrôle du juge et extinction automatique sans renouvellement explicite. L'État peut agir vite ; l'exception ne peut devenir une méthode permanente de gouvernement."),
    ]
    table = doc.add_table(rows=0, cols=1)
    for title, body in items:
        cell = table.add_row().cells[0]; set_cell_margins(cell, 105, 145, 105, 145); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2); r = p.add_run(title); style(r, 10.8, NAVY, True)
        p = cell.add_paragraph(); p.paragraph_format.space_after = Pt(0); r = p.add_run(body); style(r, 10.4)

    para(doc, "Ce qu'ils changent", style_name="Heading 1")
    para(doc, "Ces engagements ne découlent pas automatiquement du programme existant : ils créent une règle de répartition en période de pénurie, une garantie avant transformation d'un service et une limite matérielle au pouvoir exceptionnel.", after=8)

    para(doc, "Ce que nous demandons", style_name="Heading 1")
    para(doc, "Les intégrer au programme 2027 comme ajouts distincts — ou les écarter séparément. Formulations précises, sources, limites juridiques et budgétaires, conséquences territoriales et autres garanties sont disponibles pour instruction.", after=5)
    para(doc, "Contribution citoyenne : propositions ouvertes à la discussion, à la modification ou au rejet séparé.", 9.6, MUTED, italic=True, after=0)

    footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT; r = footer.add_run("Contribution citoyenne · LFI 2027"); style(r, 8.7, MUTED, True)
    OUT.parent.mkdir(exist_ok=True); doc.save(OUT); print(OUT)

if __name__ == "__main__": main()
